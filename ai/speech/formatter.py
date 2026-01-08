"""
Format transcribed text into structured documents
"""
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from config import TRANSCRIPTS_DIR

class TextFormatter:
    """Format transcribed text into structured documents"""
    
    def __init__(self):
        """Initialize formatter"""
        pass
    
    def format_as_structured_text(self, text: str, segments: List[Dict] = None) -> str:
        """
        Format text with basic structure
        
        Args:
            text: Transcribed text
            segments: Optional timestamp segments
            
        Returns:
            Formatted text with basic structure
        """
        # Basic formatting without LLM (for now)
        # Split into paragraphs based on pauses (sentences)
        sentences = text.split('. ')
        
        formatted_lines = []
        formatted_lines.append("## Lecture Transcript\n")
        
        # Group sentences into paragraphs (every 3-4 sentences)
        paragraph = []
        for i, sentence in enumerate(sentences):
            sentence = sentence.strip()
            if not sentence:
                continue
                
            paragraph.append(sentence)
            
            # Create paragraph break every 3-4 sentences
            if len(paragraph) >= 3 or i == len(sentences) - 1:
                formatted_lines.append('. '.join(paragraph) + '.\n')
                paragraph = []
        
        return '\n'.join(formatted_lines)
    
    def format_with_timestamps(self, segments: List[Dict]) -> str:
        """
        Format text with timestamps for each segment
        
        Args:
            segments: List of segments with timestamps
            
        Returns:
            Formatted text with timestamps
        """
        formatted = []
        formatted.append("## Lecture Transcript (with timestamps)\n")
        
        for seg in segments:
            start_time = self._format_time(seg.get('start', 0))
            end_time = self._format_time(seg.get('end', 0))
            text = seg.get('text', '').strip()
            
            formatted.append(f"**[{start_time} - {end_time}]**")
            formatted.append(f"{text}\n")
        
        return '\n'.join(formatted)
    
    def _format_time(self, seconds: float) -> str:
        """Convert seconds to MM:SS format"""
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes:02d}:{secs:02d}"
    
    def export_to_docx(
        self,
        text: str,
        filename: str,
        title: str = "Lecture Transcript",
        segments: List[Dict] = None
    ) -> str:
        """
        Export formatted text to DOCX document
        
        Args:
            text: Formatted text
            filename: Output filename
            title: Document title
            segments: Optional timestamp segments
            
        Returns:
            Path to saved document
        """
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=1)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=2)
            elif line.startswith('**[') and ']**' in line:
                # Timestamp line
                doc.add_paragraph(line, style='Intense Quote')
            else:
                doc.add_paragraph(line)
        
        # Save document
        output_path = TRANSCRIPTS_DIR / f"{filename}.docx"
        doc.save(output_path)
        
        print(f"📄 Document saved: {output_path}")
        return str(output_path)
    
    def export_to_markdown(
        self,
        text: str,
        filename: str,
        title: str = "Lecture Transcript"
    ) -> str:
        """
        Export formatted text to Markdown
        
        Args:
            text: Formatted text
            filename: Output filename
            title: Document title
            
        Returns:
            Path to saved document
        """
        output_path = TRANSCRIPTS_DIR / f"{filename}.md"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"# {title}\n\n")
            f.write(text)
        
        print(f"📝 Markdown saved: {output_path}")
        return str(output_path)


# Optional: Advanced formatter with LLM (if you want to add later)
class AdvancedTextFormatter(TextFormatter):
    """Format with LLM for better structure detection"""
    
    def __init__(self):
        """Initialize with LLM"""
        super().__init__()
        try:
            from rag.generator import get_generator
            self.generator = get_generator()
            self.use_llm = True
        except Exception as e:
            print(f"⚠️ LLM not available for formatting: {e}")
            self.use_llm = False
    
    def format_as_structured_text(self, text: str, segments: List[Dict] = None) -> str:
        """Format with LLM if available, otherwise use basic formatting"""
        
        if not self.use_llm:
            return super().format_as_structured_text(text, segments)
        
        # Use LLM to detect structure
        prompt = f"""Format this lecture transcript with headings and structure.

Rules:
1. Add main headings (##) for major topics
2. Add subheadings (###) for subtopics  
3. Keep original text
4. Organize into paragraphs

Transcript:
{text[:2000]}

Formatted:"""
        
        try:
            context = ""  # No context needed
            formatted = self.generator.generate_response(prompt, context)
            return formatted
        except Exception as e:
            print(f"⚠️ LLM formatting failed: {e}. Using basic formatting.")
            return super().format_as_structured_text(text, segments)
